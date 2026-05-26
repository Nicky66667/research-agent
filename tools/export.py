import io
from datetime import datetime
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import cm
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.enums import TA_CENTER


def _parse_markdown_lines(md_text: str) -> list[tuple[str, str]]:
    """
    Parse Markdown text into a list of (type, content) tuples for later formatting.
    This shared parser means all three export formats use the same parsing logic —
    only the rendering layer differs.

    Recognized types: h1 / h2 / h3 / bullet / plain / blank
    """
    result = []

    for line in md_text.splitlines():
        stripped = line.strip()
        if stripped.startswith("### "):           # h3 must be checked before h2/h1
            result.append(("h3", stripped[4:]))   # slice off the "### " prefix
        elif stripped.startswith("## "):
            result.append(("h2", stripped[3:]))
        elif stripped.startswith("# "):
            result.append(("h1", stripped[2:]))
        elif stripped.startswith(("- ", "* ")):   # both common bullet styles
            result.append(("bullet", stripped[2:]))
        elif stripped == "":
            result.append(("blank", ""))          # preserve paragraph spacing
        else:
            result.append(("plain", stripped))    # regular body text

    return result


def export_txt(report_md: str, query: str = "") -> bytes:
    """
    Export the report as plain UTF-8 text.
    Markdown syntax is stripped and replaced with simple ASCII formatting
    (uppercase headings, dashes, bullet dots) so the file is readable
    in any text editor without rendering.
    """
    # Build a human-readable file header
    header = "Research Report\n"
    if query:
        header += f"Query: {query}\n"
    header += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}\n"
    header += "=" * 60 + "\n\n"

    lines = []
    for t, content in _parse_markdown_lines(report_md):
        if t in ("h1", "h2"):
            # Capitalise headings and underline them with dashes for visual weight
            lines.append(f"\n{content.upper()}\n" + "-" * len(content))
        elif t == "h3":
            lines.append(f"\n{content}")          # sub-headings get a blank line above
        elif t == "bullet":
            lines.append(f"  • {content}")        # indent bullet points
        elif t == "blank":
            lines.append("")                      # keep paragraph breaks
        else:
            lines.append(content)                 # plain body text, no change

    body = "\n".join(lines)
    return (header + body).encode("utf-8")        # encode to bytes for st.download_button


def export_docx(report_md: str, query: str = "") -> bytes:
    """
    Export the report as a Word document (.docx) using python-docx.
    python-docx generates the file in pure Python — no Word installation needed,
    which means this works on any OS including the HF Spaces Docker container.
    """
    doc = Document()

    # ── Title block ──────────────────────────────────────────────
    title = doc.add_heading("Research Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if query:
        # Show the original query in italic grey text below the title
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Query: {query}")
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Timestamp in small print, centred
    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}").font.size = Pt(9)

    doc.add_paragraph()  # visual spacer before body content

    # ── Body content ─────────────────────────────────────────────
    for t, content in _parse_markdown_lines(report_md):
        if t == "h1":
            doc.add_heading(content, level=1)
        elif t == "h2":
            doc.add_heading(content, level=2)
        elif t == "h3":
            doc.add_heading(content, level=3)
        elif t == "bullet":
            # "List Bullet" is a built-in Word style — renders as a proper bullet point
            doc.add_paragraph(content, style="List Bullet")
        elif t == "blank":
            pass  # docx paragraph spacing already provides visual separation
        else:
            doc.add_paragraph(content)

    # Save to an in-memory buffer instead of disk, then return raw bytes
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def export_pdf(report_md: str, query: str = "") -> bytes:
    """
    Export the report as a PDF using reportlab.
    reportlab is pure Python and needs no system-level dependencies (unlike weasyprint,
    which requires Cairo/Pango), making it the safest choice for Docker deployments.

    Note: reportlab's default fonts don't support CJK characters.
    If the report contains Chinese/Japanese/Korean text, register a TTF font first:
        from reportlab.pdfbase import pdfmetrics
        from reportlab.pdfbase.ttfonts import TTFont
        pdfmetrics.registerFont(TTFont('NotoSans', '/path/to/NotoSansSC-Regular.ttf'))
    then set fontName='NotoSans' in each ParagraphStyle below.
    """
    buf = io.BytesIO()

    # SimpleDocTemplate handles page layout, margins, and pagination automatically
    doc = SimpleDocTemplate(
        buf,
        pagesize=A4,
        leftMargin=2.5 * cm,
        rightMargin=2.5 * cm,
        topMargin=2.5 * cm,
        bottomMargin=2.5 * cm,
    )

    styles = getSampleStyleSheet()  # reportlab's built-in base styles

    # ── Custom paragraph styles ───────────────────────────────────
    style_title = ParagraphStyle(
        "ReportTitle",
        parent=styles["Title"],
        fontSize=20,
        spaceAfter=6,
        alignment=TA_CENTER,
    )
    style_meta = ParagraphStyle(
        "Meta",
        parent=styles["Normal"],
        fontSize=9,
        textColor=colors.grey,
        spaceAfter=20,
        alignment=TA_CENTER,
    )
    style_h2 = ParagraphStyle(
        "H2",
        parent=styles["Heading2"],
        fontSize=13,
        spaceBefore=14,
        spaceAfter=4,
        textColor=colors.HexColor("#1a1a2e"),
    )
    style_h3 = ParagraphStyle(
        "H3",
        parent=styles["Heading3"],
        fontSize=11,
        spaceBefore=8,
        spaceAfter=2,
    )
    style_body = ParagraphStyle(
        "Body",
        parent=styles["Normal"],
        fontSize=10,
        leading=15,   # line height; 1.5x font size for comfortable reading
        spaceAfter=4,
    )
    style_bullet = ParagraphStyle(
        "Bullet",
        parent=styles["Normal"],
        fontSize=10,
        leading=14,
        leftIndent=16,   # indent the whole bullet line
        bulletIndent=6,  # indent the bullet symbol itself
        spaceAfter=2,
    )

    # ── Build the story (reportlab's term for a list of flowable elements) ──
    story = []

    # Title and metadata at the top of the document
    story.append(Paragraph("AI Research Report", style_title))
    meta_text = ""
    if query:
        meta_text += f"Query: {query}&nbsp;&nbsp;|&nbsp;&nbsp;"  # &nbsp; = non-breaking space in reportlab
    meta_text += f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    story.append(Paragraph(meta_text, style_meta))

    # Body content
    for t, content in _parse_markdown_lines(report_md):
        # Escape HTML special characters — reportlab uses XML-like markup internally
        content_safe = content.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        if t in ("h1", "h2"):
            story.append(Paragraph(content_safe, style_h2))
        elif t == "h3":
            story.append(Paragraph(content_safe, style_h3))
        elif t == "bullet":
            story.append(Paragraph(f"• {content_safe}", style_bullet))
        elif t == "blank":
            story.append(Spacer(1, 6))  # 6-point vertical gap between paragraphs
        else:
            story.append(Paragraph(content_safe, style_body))

    # Render all flowables into the PDF and write to buffer
    doc.build(story)
    return buf.getvalue()  # return raw bytes for st.download_button